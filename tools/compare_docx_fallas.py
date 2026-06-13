from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document

DOCX_PATH = Path(r"C:\Users\marcb\Desktop\fallas360\Fallas_360_base_fallas_2026_actualizada.docx")
OUT_DIR = Path("runtime")
OUT_DIR.mkdir(exist_ok=True)


def clean(value: str) -> str:
    return " ".join((value or "").replace("\xa0", " ").split()).strip()


doc = Document(str(DOCX_PATH))

if len(doc.tables) < 3:
    raise SystemExit("El documento no contiene las tablas esperadas.")

summary_table = doc.tables[0]
prize_table = doc.tables[1]
consolidated_table = doc.tables[2]

summary = {}
for row in summary_table.rows[1:]:
    cells = [clean(cell.text) for cell in row.cells]
    if len(cells) >= 2:
        summary[cells[0]] = cells[1]

prizes = []
headers = [clean(cell.text) for cell in prize_table.rows[0].cells]
for row in prize_table.rows[1:]:
    cells = [clean(cell.text) for cell in row.cells]
    if any(cells):
        prizes.append(dict(zip(headers, cells)))

headers = [clean(cell.text) for cell in consolidated_table.rows[0].cells]
rows = []
for row in consolidated_table.rows[1:]:
    cells = [clean(cell.text) for cell in row.cells]
    if any(cells):
        rows.append(dict(zip(headers, cells)))

main_updates = []
infant_updates = []
for row in rows:
    main_id = clean(row.get("ID principal/app", ""))
    if main_id.isdigit():
        main_updates.append(
            {
                "id": int(main_id),
                "name": clean(row.get("Nombre falla", "")),
                "category": clean(row.get("Tipo app", "")),
                "section_name": clean(row.get("Sección principal/app", "")),
                "neighborhood": clean(row.get("Barrio/Zona", "")),
                "cross": clean(row.get("Cruce", "")),
            }
        )

    infant_id = clean(row.get("ID infantil", ""))
    if infant_id.isdigit():
        infant_updates.append(
            {
                "id": int(infant_id),
                "nombre": clean(row.get("Nombre falla", "")),
                "seccion_label": clean(row.get("Sección infantil", "")),
                "artista": clean(row.get("Artista infantil", "")),
                "presupuesto_formateado": clean(row.get("Presupuesto infantil", "")),
                "ciudad": clean(row.get("Ciudad", "")),
                "cross": clean(row.get("Cruce", "")),
            }
        )

payload = {
    "summary": summary,
    "prizes": prizes,
    "counts": {
        "consolidated_rows": len(rows),
        "main_rows_with_id": len(main_updates),
        "infant_rows_with_id": len(infant_updates),
    },
}

(OUT_DIR / "docx_fallas_summary.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

with (OUT_DIR / "docx_fallas_main_updates.csv").open("w", encoding="utf-8", newline="") as handle:
    writer = csv.DictWriter(handle, fieldnames=["id", "name", "category", "section_name", "neighborhood", "cross"])
    writer.writeheader()
    writer.writerows(main_updates)

with (OUT_DIR / "docx_fallas_infant_updates.csv").open("w", encoding="utf-8", newline="") as handle:
    writer = csv.DictWriter(handle, fieldnames=["id", "nombre", "seccion_label", "artista", "presupuesto_formateado", "ciudad", "cross"])
    writer.writeheader()
    writer.writerows(infant_updates)

print(json.dumps(payload, ensure_ascii=False, indent=2))
